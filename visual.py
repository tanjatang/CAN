"""
Training script. Should be pretty adaptable to whatever.
"""
import os
os.environ["CUDA_VISIBLE_DEVICES"] = "5"
import argparse
import multiprocessing
from allennlp.common.params import Params
from allennlp.training.optimizers import Optimizer
from torch.nn import DataParallel
from dataloaders.vcr_new1 import VCR, VCRLoader
import logging
logging.basicConfig(format='%(asctime)s - %(levelname)s - %(name)s - %(message)s', level=logging.DEBUG)
from utils.pytorch_misc import time_batch, save_checkpoint, clip_grad_norm, restore_checkpoint, print_para, restore_best_checkpoint
from docx import Document
from docx.shared import Inches
# This is needed to make the imports work
from allennlp.models import Model
import models
import warnings
warnings.filterwarnings("ignore")

parser = argparse.ArgumentParser(description='train')
parser.add_argument(
    '-params',
    dest='params',
    help='Params location',
    type=str,
)
parser.add_argument(
    '-rationale',
    action="store_true",
    help='use rationale',
)
parser.add_argument(
    '-folder',
    dest='folder',
    help='folder location',
    type=str,
)
parser.add_argument(
    '-no_tqdm',
    dest='no_tqdm',
    action='store_true',
)

args = parser.parse_args()
print(str(args.params))
params = Params.from_file(args.params)
# ipdb.set_trace()
train, val, test = VCR.splits(mode='rationale' if args.rationale else 'answer',
                              embs_to_load=params['dataset_reader'].get('embs', 'bert_da'),
                              only_use_relevant_dets=params['dataset_reader'].get('only_use_relevant_dets', True))

#q确定模式为train，val还是test，从bert下载相应的数据
NUM_GPUS = 1#torch.cuda.device_count()
NUM_CPUS = 1#multiprocessing.cpu_count()
if NUM_GPUS == 0:
    raise ValueError("you need gpus!")


num_workers = 1#(4 * NUM_GPUS if NUM_CPUS == 32 else 2*NUM_GPUS)-1

print(f"Using {num_workers} workers out of {NUM_CPUS} possible", flush=True)
loader_params = {'batch_size': 1// NUM_GPUS, 'num_gpus':NUM_GPUS, 'num_workers':num_workers} #96
train_loader = VCRLoader.from_dataset(train, **loader_params)
val_loader = VCRLoader.from_dataset(val, **loader_params)
# test_loader = VCRLoader.from_dataset(test, **loader_params)


################################################################################

def _to_gpu(td):
    if NUM_GPUS > 1:
        return td
    for k in td:

        if k != 'metadata':

            td[k] = {k2: v.cuda(non_blocking=True) for k2, v in td[k].items()} if isinstance(td[k], dict) else td[k].cuda(
                non_blocking=True) #transfer tensor to cuda
    return td

model = Model.from_params(vocab=train.vocab, params=params['model'])


model = DataParallel(model).cuda() if NUM_GPUS > 1 else model.cuda()
optimizer = Optimizer.from_params([x for x in model.named_parameters() if x[1].requires_grad],
                                  params['trainer']['optimizer'])#在default里面，用Adam优化，

if __name__ == '__main__':
    # with open('vs.docx','wr') as f:
    #     doc = f
    doc = Document()
    restore_best_checkpoint(model, args.folder)
    model.eval()
    for i ,j in enumerate(val_loader):
        if i == 500:
            doc.save('rvs1.docx')
            break
        collect = j
        batch = _to_gpu(collect)
        output_dict = model(**batch)
        doc.add_paragraph(str(i)+':    ')
        # doc.add_paragraph( str(output_dict['label_probs']))
        doc.add_paragraph(str(output_dict['label_probs'].argmax()))

        # print("metadata: ", batch['metadata'])
        # print("predict answer: ", output_dict['label_probs'])
        # print("boxes:",batch['boxes'])


    # print("answer_tags: ", batch['answer_tags'])

    # img = img.squeeze(0)
    # plt.imshow(img)
    # collect = next(iter(train_loader))
    # print(len(collect),next(iter(collect)))
    # # for k,v in collect.items():
    # #     print(k, v)
